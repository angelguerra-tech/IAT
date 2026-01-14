import pandas as pd
import streamlit as st
from docx import Document
import io
from PIL import Image

# Asegurarse de que el diccionario de observaciones está inicializado
if 'observaciones_por_normativa' not in st.session_state:
    st.session_state['observaciones_por_normativa'] = {}

# --- NUEVO: estados para flujo reporte/dictamen (mínimamente invasivo) ---
if "reporte_docx_bytes" not in st.session_state:
    st.session_state["reporte_docx_bytes"] = None
if "reporte_generado" not in st.session_state:
    st.session_state["reporte_generado"] = False
if "reporte_descargado" not in st.session_state:
    st.session_state["reporte_descargado"] = False
if "reporte_tiene_no_cumple" not in st.session_state:
    st.session_state["reporte_tiene_no_cumple"] = False
if "dictamen_observaciones" not in st.session_state:
    st.session_state["dictamen_observaciones"] = ""

# Funciones para manejar observaciones
def agregar_observacion(norma, requisito, cumplimiento, observacion):
    if norma not in st.session_state['observaciones_por_normativa']:
        st.session_state['observaciones_por_normativa'][norma] = {}
    st.session_state['observaciones_por_normativa'][norma][requisito] = {
        'cumplimiento': cumplimiento,
        'observacion': observacion
    }

def eliminar_observacion(norma, requisito):
    if norma in st.session_state['observaciones_por_normativa']:
        if requisito in st.session_state['observaciones_por_normativa'][norma]:
            del st.session_state['observaciones_por_normativa'][norma][requisito]

# Función para cambiar de normativa
def cambiar_normativa(norma):
    st.session_state['current_norma'] = norma


def clear_input_states():
    keys_to_delete = []
    for key in st.session_state.keys():
        if key.startswith('indice_') or key.startswith('input_'):
            keys_to_delete.append(key)
    if 'observaciones_por_normativa' in st.session_state:
        keys_to_delete.append('observaciones_por_normativa')
    for key in keys_to_delete:
        del st.session_state[key]

def clear_specific_keys():
    text_area_keys = [key for key in st.session_state.keys() if key.startswith('input_') or key.startswith('obs_')]
    for key in text_area_keys:
        st.session_state[key] = ""

    selectbox_keys = [key for key in st.session_state.keys() if key.startswith('indice_')]
    for key in selectbox_keys:
        st.session_state[key] = 0


st.set_page_config(layout="wide")


st.markdown("""
<style>
[data-testid="stSidebarContent"] { background: #F4F7FD; }
button.st-emotion-cache-1mcbg9u e16zdaao0 { background-color: #FFFFFF !important; }
div.st-emotion-cache-1jicfl2 { background-color: white; }
[data-testid="stBaseButton-secondary"] {
    background: #436ab2 !important;
    color: white !important;
    border: 2px solid #436ab2 !important;
}
[data-testid="stMarkdownContainer"] { display: block; width: 100%; text-align: left; }
button[data-testid="stBaseButton-secondary"]:hover {
    background-color: #FFFFFF !important;
    color: #436ab2 !important;
    border-color: #436ab2 !important;
}
div[data-testid="stButton"] > button { display: block; margin: 0 auto; }
div[data-testid="stDownloadButton"] > button { display: block; margin: 0 auto; }             
[data-testid="stBaseButton-primary"] {
    background-color: #81d4fa;
    color: black;
    border: 2px solid #81d4fa;
    text-align: center;
}
[data-testid="bstBaseButton-primary"]:hover {
    background-color: #4fc3f7;
    color: black;
    border-color: #29b6f6;
    text-align: center;
}
[data-testid="stLinkButton"] { width: 100%; text-align: right; }
[data-testid="stLinkButton"]>a {
    background-color: #4dd0e1;
    color: white;
    border: 2px solid #4dd0e1;
    padding: 8px 16px;
    border-radius: 5px;
}
[data-testid="stLinkButton"]>a:hover {
    background-color: #26c6da;
    border-color: #00acc1;
}
</style>
""", unsafe_allow_html=True)

observaciones_no_cumple = []


def load_data(ruta_excel):
    return pd.read_excel(ruta_excel, sheet_name=None, keep_default_na=False, na_values='')

ruta_excel = 'Matriz.xlsx'
matriz = load_data(ruta_excel)

categorias_df = matriz['Vinculación de CA']
reglamentos_df = matriz['Reglamentos Aplicables']
requisitos_df = matriz['REQUISITOS']

if 'observaciones_por_normativa' not in st.session_state:
    st.session_state['observaciones_por_normativa'] = {}
    
if 'current_norma' not in st.session_state:
    st.session_state['current_norma'] = None


logo = Image.open("logo_3.png")
st.sidebar.image(logo, width=150, use_container_width=True)

space1, col1, col2, space = st.sidebar.columns([4,2,2,4], vertical_alignment="center")

with space1:
    if st.button("❌", type='primary', help="Cerrar Sesión"):
        st.session_state['salida'] = True
        st.switch_page("evaluacion_alimentos.py")
with col1:
    if st.button("🢀", type='primary'):
        clear_specific_keys()
        st.switch_page("pages/CATEGORIAS.py")
with col2:
    if st.button("🏚️", type='primary'):
        clear_specific_keys()
        st.switch_page("evaluacion_alimentos.py")


st.sidebar.title("Regulación aplicable")

categoria_seleccionada = st.session_state.get('categoria_seleccionada', None)

if 'categoria_seleccionada' not in st.session_state:
    st.session_state['categoria_seleccionada'] = None

if 'last_categoria_seleccionada' not in st.session_state:
    st.session_state['last_categoria_seleccionada'] = None

col1, col2, col3, col4 = st.columns([0.5,0.5,8, 2], vertical_alignment="center")

with col3: 
    if 'categoria_seleccionada' in st.session_state:
        categoria_seleccionada = st.session_state['categoria_seleccionada']
        subgrupo = st.session_state['Subgrupo_rtca']
        Descrip_sub_rtca = st.session_state['Descriptor_subcategoria_rtca']

        st.markdown(f"""
        <div style="text-align: center; font-size: 25px;">
            <strong><em>{categoria_seleccionada}</em></strong>
        </div>
        <div style="text-align: left; font-size: 16px;">
            <strong><em>Subgrupo RTCA: {Descrip_sub_rtca}</em></strong>
        </div>
        """, unsafe_allow_html=True)

st.write("")
st.write("")
st.markdown("""
<div style="text-align: left; font-size: 15px; color: #636280">
    <strong>En esta sección se presenta la regulación aplicable correspondiente a la categoría de alimento en evaluación.<br><br>
    Para acceder a los requisitos correspondientes de cada regulación, hacer click en el reglamento a verificar. <br><br>
    Para evaluar un requisito en específico debe hacer click en ver requisito, esto lo direccionará a la sección específica del reglamento 
            a evaluar.<br><br>
    Una vez verificado el cumplimiento del requisito correspondiente debe seleccionar en el menú desplegable según corresponda 
            (cumple, no cumple, no aplica). En caso de incumplimiento es mandatorio emitir observaciones de acuerdo a la naturaleza 
            del hallazgo.<br><br>
    Al finalizar la verificación de los requisitos correspondientes a la regulación aplicable, hacer click en generar reporte y luego hacer 
            click en descargar. Esto guardará automáticamente en la bandeja de descargas las observaciones resultantes de la observación.<br><br>
    Las observaciones resultantes del IAT deberán colocarse en la sección de observaciones de SISAM.
</strong>
</div>
""", unsafe_allow_html=True)

if st.session_state['categoria_seleccionada'] != st.session_state['last_categoria_seleccionada']:
    st.session_state['current_norma'] = None
    st.session_state['last_categoria_seleccionada'] = st.session_state['categoria_seleccionada']


if categoria_seleccionada:
    normas_aplicables = categorias_df[categorias_df['Subcategoria'] == categoria_seleccionada].iloc[0, 8:].dropna().tolist()


if 'current_norma' not in st.session_state:
    st.session_state['current_norma'] = None


for norma in normas_aplicables:
    if st.sidebar.button(norma, use_container_width=True):
        cambiar_normativa(norma)

st.markdown("---")

if 'show_selectbox' not in st.session_state:
    st.session_state.show_selectbox = False

if 'selected_fortification' not in st.session_state:
    st.session_state.selected_fortification = None

cumplimiento_guardado = None
observacion_guardada = None


if 'current_norma' in st.session_state and st.session_state['current_norma']:

    norma = st.session_state['current_norma'].strip()

    st.markdown(f"""
    <div style="text-align: center; font-size: 22px; color:#005662;">
        <strong><em>{norma}</em></strong>
    </div>
    """, unsafe_allow_html=True)

    reglamentos_df['REGLAMENTO'] = reglamentos_df['REGLAMENTO'].str.strip()
    enlace = reglamentos_df.loc[reglamentos_df['REGLAMENTO'] == norma, 'ENLACE'].iloc[0]

    if norma != "Observaciones Generales":
        st.link_button("VER REGLAMENTO", enlace)

    st.write("")
    st.write("")

    requisitos_df['Normas'] = requisitos_df['Normas'].str.strip()

    norma_forti = "RTS 67.06.01:13 Fortificación de alimentos. Especificaciones (azúcar, sal, harina de maíz nixtamalizado y pastas alimenticias)"
    if norma == norma_forti:
        st.session_state.show_selectbox = True

    if st.session_state.show_selectbox:
        selected = st.sidebar.selectbox(
            "",
            ("Seleccione una opción...", "Azúcar", "Sal", "Harina de maíz nixtamalizado", "Pastas alimenticias"),
            key="fortification_select",
            placeholder="Seleccione la sección aplicable:",
        )
        st.session_state.selected_fortification = selected
        requisitos = requisitos_df[(requisitos_df['Normas'] == norma) & (requisitos_df['INFO'] == selected)]
    else:
        requisitos = requisitos_df[requisitos_df['Normas'] == norma]

    if st.session_state.selected_fortification:
        st.markdown(f"""
        <div style="text-align: left; font-size: 18px; color:red">
            <strong>{st.session_state.selected_fortification}</strong>
        </div>
        """, unsafe_allow_html=True)

    if 'observaciones_no_cumple' not in st.session_state:
        st.session_state['observaciones_no_cumple'] = []
    

    for index, row in requisitos.iterrows():

        requisito = row['Requisito']

        indice = f'indice_{norma}_{requisito}'
        obs_key = f"obs_{norma}_{requisito}_{index}"

        if obs_key not in st.session_state:
            st.session_state[obs_key] = ""

        if indice not in st.session_state:
            st.session_state[indice] = 0

        if norma in st.session_state['observaciones_por_normativa']:
            observaciones_norma = st.session_state['observaciones_por_normativa'][norma]
            if requisito in observaciones_norma:
                observacion_guardada = observaciones_norma[requisito].get('observacion', "")
        
        especificacion = f"Sección: {row['Sección']} - {row['Requisito']}"
        obs_dictamen = f"{row['Sección']} - {row['Requisito']}"

        if norma != "Observaciones Generales":
            st.markdown(f"""
            <div style="text-align: left; font-size: 18px;">
                <strong><em>{especificacion}</em></strong>
            </div>
            """, unsafe_allow_html=True)

        observacion_guardada = st.session_state[obs_key]

        if norma == "Observaciones Generales":
            def on_text_area_change(key):
                st.session_state[key] = st.session_state[f"input_{key}"]

            input_key = f"input_{obs_key}"
            observacion = st.text_area("", value=observacion_guardada, key=input_key, on_change=on_text_area_change, args=(obs_key,))
            agregar_observacion(norma, requisito="General", cumplimiento="No cumple", observacion=observacion)
            break

        if norma == "RTCA 67.04.54:18 Alimentos y bebidas procesadas. Aditivos alimentarios" and not pd.isna(row['LINK']) and (row['Sección'] == "Cuadro 2" or row['Sección'] == "Cuadro 3"):
            st.link_button("Norma Codex", row['LINK'])
        else:
            st.link_button("Ver requisito", row['LINK'])

        if st.session_state.selected_fortification and norma == norma_forti:
            opciones_cumple = ('No aplica', 'No cumple', 'Cumple')
        else:
            opciones_cumple = ('Cumple', 'No cumple', 'No aplica')

        cumplimiento = st.selectbox("", opciones_cumple, index=st.session_state[indice], key=f"cumple_{norma}_{index}")

        def on_text_area_change(key):
            st.session_state[key] = st.session_state[f"input_{key}"]

        is_disabled = cumplimiento != "No cumple"
        
        input_key = f"input_{obs_key}"
        observacion = st.text_area("Observaciones", value=observacion_guardada, key=input_key, on_change=on_text_area_change, args=(obs_key,), disabled=is_disabled)
        
        st.session_state[indice] = opciones_cumple.index(cumplimiento)

        eliminar_observacion(norma, obs_dictamen)
        agregar_observacion(norma, obs_dictamen, cumplimiento, observacion)

        if norma == "RTS 67.07.01:22 Mezcla de Crema (Nata) con Aceite o Grasa Vegetal Comestible. Especificaciones" and categoria_seleccionada == "1.4.4 Productos análogos a la nata (crema)":
            break

    if 'current_norma' in st.session_state and st.session_state['current_norma'] != norma:
        st.session_state.show_selectbox = False
        st.session_state.selected_fortification = None


def generar_reporte():
    tiene_no_cumple = any(
        obs_data['cumplimiento'] == "No cumple"
        for obs_list in st.session_state['observaciones_por_normativa'].values()
        for obs_data in obs_list.values()
    )

    template_path = "Dictamen_desfavorable_2.docx" if tiene_no_cumple else "Dictamen_favorable.docx"
    doc = Document(template_path)

    observaciones_texto = ""

    if tiene_no_cumple:
        contador = 1
        for norma, requisitos in st.session_state['observaciones_por_normativa'].items():
            for requisito, obs in requisitos.items():
                if obs['cumplimiento'] == "No cumple":
                    if norma == "Observaciones Generales":
                        observaciones_texto += f"{contador}. Incumplimiento en: {obs['observacion']}.\n"
                    else:
                        observaciones_texto += (
                            f"{contador}. Incumplimiento con el numeral {requisito} del Reglamento {norma} "
                            f"en cuanto a: {obs['observacion']}.\n"
                        )
                    contador += 1
        
        placeholder = "{Observaciones}"
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, observaciones_texto.strip())
        
        cumplimientos_texto = ""
        normas_con_cumplimientos = {}
        for norma, requisitos in st.session_state['observaciones_por_normativa'].items():
            if norma != "Observaciones Generales":
                for requisito, obs in requisitos.items():
                    if obs['cumplimiento'] == "Cumple":
                        normas_con_cumplimientos.setdefault(norma, []).append(requisito)

        for norma, requisitos_lista in normas_con_cumplimientos.items():
            cumplimientos_texto += f"• {norma}\n"
            for requisito in requisitos_lista:
                cumplimientos_texto += f"  - {requisito}\n"
        
        placeholder_cumplimientos = "{Cumplimientos}"
        for paragraph in doc.paragraphs:
            if placeholder_cumplimientos in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder_cumplimientos, cumplimientos_texto.strip())
        
        inaplicables_texto = ""
        normas_con_inaplicables = {}
        for norma, requisitos in st.session_state['observaciones_por_normativa'].items():
            if norma != "Observaciones Generales":
                for requisito, obs in requisitos.items():
                    if obs['cumplimiento'] == "No aplica":
                        normas_con_inaplicables.setdefault(norma, []).append(requisito)

        for norma, requisitos_lista in normas_con_inaplicables.items():
            inaplicables_texto += f"• {norma}\n"
            for requisito in requisitos_lista:
                inaplicables_texto += f"   - {requisito}\n"

        placeholder_inaplicables = "{Inaplicables}"
        for paragraph in doc.paragraphs:
            if placeholder_inaplicables in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder_inaplicables, inaplicables_texto.strip())

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    reporte_bytes = output.getvalue()

    # --- NUEVO: guardar estado para dictamen ---
    st.session_state["reporte_docx_bytes"] = reporte_bytes
    st.session_state["reporte_generado"] = True
    st.session_state["reporte_descargado"] = False
    st.session_state["reporte_tiene_no_cumple"] = tiene_no_cumple
    st.session_state["dictamen_observaciones"] = observaciones_texto.strip() if tiene_no_cumple else ""

    return reporte_bytes

def construir_cumplimientos_para_dictamen():
    EXCLUIR = {
        "RTCA 67.01.31:20 Alimentos procesados.Procedimiento para el otorgamiento, renovación y modificación del registro sanitario",
        "RTCA 67.04.50:17 Alimentos. Criterios microbiológicos para la inocuidad de alimentos",
    }

    normas = []
    for norma, reqs in st.session_state.get("observaciones_por_normativa", {}).items():
        if norma == "Observaciones Generales":
            continue
        norma_limpia = str(norma).strip()
        if norma_limpia in EXCLUIR:
            continue

        # Solo incluir si NO tiene "No cumple" en esa norma
        if all(obs.get("cumplimiento") != "No cumple" for obs in reqs.values()):
            normas.append(norma_limpia)

    # Texto tipo viñetas para {Cumplimientos}
    return "\n".join([f"• {n}" for n in normas])


def _marcar_descarga_reporte():
    st.session_state["reporte_descargado"] = True


# 1) Generar reporte
if st.sidebar.button('Generar reporte', type='primary'):
    generar_reporte()

# 2) Descargar (visible si ya hay bytes)
if st.session_state.get("reporte_docx_bytes"):
    st.sidebar.download_button(
        "Descargar",
        data=st.session_state["reporte_docx_bytes"],
        file_name="Reporte.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type='primary',
        on_click=_marcar_descarga_reporte,
        key="download_reporte",
    )

# 3) Nuevo paso: Generar dictamen (solo si hubo "No cumple")
if st.session_state.get("reporte_descargado", False):
    if st.session_state.get("reporte_tiene_no_cumple", False):
        st.sidebar.info("Se detectaron incumplimientos. Puedes emitir el dictamen.")
        if st.sidebar.button("Generar dictamen", type="secondary"):
            st.switch_page("pages/01_Dictamen_no_conformidad.py")
    else:
        st.sidebar.success("Todos los requisitos cumplen. (Dictamen no habilitado por ahora).")
