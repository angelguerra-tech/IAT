from pathlib import Path
from docx import Document
import streamlit as st
import pandas as pd
import io
import re
import unicodedata
from datetime import date

BASE_DIR = Path(__file__).resolve().parents[1]  # sube de /pages a la raíz del proyecto
RUTA_EXCEL = BASE_DIR / "Matriz.xlsx"
PLANTILLA_DICTAMEN = BASE_DIR / "Dictamen_conformidad.docx"


def _norm(s: str) -> str:
    s = str(s).strip().lower()
    s = "".join(c for c in unicodedata.normalize("NFD", s) if unicodedata.category(c) != "Mn")
    s = s.replace("_", " ").replace("-", " ")
    s = re.sub(r"\s+", " ", s)
    return s


def _load_catalogos():
    df = pd.read_excel(RUTA_EXCEL, sheet_name="Catalogos", keep_default_na=False)

    colmap = {_norm(c): c for c in df.columns}

    def pick(*cands):
        for c in cands:
            c = _norm(c)
            if c in colmap:
                return colmap[c]
        return None

    col_titulo = pick("titulo", "título")
    col_nombre = pick("nombre del analista", "nombre analista", "analista", "nombre")
    col_cargo = pick("cargo", "puesto")

    if not col_titulo or not col_nombre or not col_cargo:
        st.error(
            "En la hoja 'Catalogos' faltan columnas esperadas. "
            f"Detectadas: {list(df.columns)}"
        )
        st.stop()

    titulos = [x for x in df[col_titulo].astype(str).tolist() if str(x).strip()]
    nombres = [x for x in df[col_nombre].astype(str).tolist() if str(x).strip()]
    cargos = [x for x in df[col_cargo].astype(str).tolist() if str(x).strip()]

    return titulos, nombres, cargos


def _replace_in_paragraph(paragraph, replacements: dict):
    full = "".join(run.text for run in paragraph.runs)
    new = full
    for k, v in replacements.items():
        new = new.replace(k, v)

    if new != full:
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = new
        else:
            paragraph.add_run(new)


def _replace_everywhere(doc: Document, replacements: dict):
    for p in doc.paragraphs:
        _replace_in_paragraph(p, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, replacements)

    # por si el template usa header/footer
    for section in doc.sections:
        for p in section.header.paragraphs:
            _replace_in_paragraph(p, replacements)
        for p in section.footer.paragraphs:
            _replace_in_paragraph(p, replacements)


st.set_page_config(layout="wide")
st.title("Generar dictamen (Conformidad)")

# 1) Trazabilidad: debe existir descarga previa del reporte
if not st.session_state.get("reporte_descargado", False):
    st.warning("Primero debes generar y descargar el reporte para mantener trazabilidad.")
    st.stop()

# 2) Este dictamen SOLO aplica cuando NO hay incumplimientos
if st.session_state.get("reporte_tiene_no_cumple", False):
    st.info("Se detectaron incumplimientos. Este dictamen es solo para cumplimiento completo.")
    st.stop()

# 3) Validar archivos base (errores claros)
if not RUTA_EXCEL.exists():
    st.error(f"No se encontró el archivo Excel: {RUTA_EXCEL}")
    st.stop()

if not PLANTILLA_DICTAMEN.exists():
    st.error(
        "No se encontró la plantilla del dictamen de conformidad. "
        f"Verifica que exista en la raíz del proyecto con este nombre exacto: {PLANTILLA_DICTAMEN.name}"
    )
    st.stop()

titulos, nombres, cargos = _load_catalogos()

c1, c2 = st.columns(2)
with c1:
    fecha = st.date_input("Fecha", value=date.today())
with c2:
    solicitud = st.text_input("Solicitud", value=st.session_state.get("solicitud", ""))

nombre_producto = st.text_input(
    "Nombre del producto",
    value=st.session_state.get("nombre_producto", "")
)

d1, d2, d3 = st.columns(3)
with d1:
    titulo_sel = st.selectbox("Título", options=titulos)
with d2:
    nombre_sel = st.selectbox("Nombre analista", options=nombres)
with d3:
    cargo_sel = st.selectbox("Cargo", options=cargos)

st.write("")

if st.button("Generar dictamen", type="primary"):
    # Este texto se construye en Requisitos_plantilla.py usando el listado del panel izquierdo
    cumplimientos = str(st.session_state.get("dictamen_cumplimientos", "")).strip()

    # Si está vacío, NO lo metemos en el doc; mostramos error para corregir el flujo
    if not cumplimientos:
        st.error(
            "No se encontró el listado de reglamentos aplicables para {Cumplimientos}. "
            "Verifica que en Requisitos_plantilla.py se esté guardando session_state['dictamen_cumplimientos'] "
            "cuando NO hay incumplimientos."
        )
        st.stop()

    replacements = {
        "{Fecha}": fecha.strftime("%d/%m/%Y"),
        "{Solicitud}": str(solicitud).strip(),
        "{Nombre del alimento}": str(nombre_producto).strip(),
        "{Título}": str(titulo_sel).strip(),
        "{Nombre}": str(nombre_sel).strip(),
        "{Cargo}": str(cargo_sel).strip(),
        "{Cumplimientos}": cumplimientos,
    }

    doc = Document(str(PLANTILLA_DICTAMEN))
    _replace_everywhere(doc, replacements)

    buff = io.BytesIO()
    doc.save(buff)
    buff.seek(0)

    st.success("Dictamen generado.")
    st.download_button(
        "Descargar dictamen",
        data=buff.getvalue(),
        file_name=f"Dictamen_conformidad_{solicitud or 's_solicitud'}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

